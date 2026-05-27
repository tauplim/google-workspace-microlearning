#!/usr/bin/env python3
"""
Convert Participant Workbook Markdown to DOCX for Google Docs Distribution
A4 size with proper pagination - each module starts on a fresh page
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()

def create_workbook_docx(output_path):
    doc = Document()

    # Set A4 page size
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ===== TITLE =====
    title = doc.add_heading('Google Workspace with AI', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Participant Workbook', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ===== HEADER INFO =====
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Name:"
    table.cell(0, 1).text = "_________________________________"
    table.cell(1, 0).text = "Date:"
    table.cell(1, 1).text = "_________________________________"
    table.cell(2, 0).text = "Session:"
    table.cell(2, 1).text = "1 of 8"

    doc.add_paragraph()

    # ===== COURSE OVERVIEW =====
    doc.add_heading('Course Overview', level=2)

    p = doc.add_paragraph()
    p.add_run('8 Modules × 30 Minutes = 4 Hours Total').bold = True

    # Module table
    mod_table = doc.add_table(rows=9, cols=3)
    mod_table.style = 'Table Grid'
    headers = ['Module', 'Topic', 'My Notes']
    for i, h in enumerate(headers):
        mod_table.cell(0, i).text = h
        mod_table.cell(0, i).paragraphs[0].runs[0].bold = True

    modules = [
        ('1', 'Foundation - Google Workspace with AI', ''),
        ('2', 'Collaboration with AI', ''),
        ('3', 'Document Automation', ''),
        ('4', 'AI-Powered Research', ''),
        ('5', 'Meetings & Calendaring', ''),
        ('6', 'Drive & File Management', ''),
        ('7', 'Automation with Gemini', ''),
        ('8', 'AI Integration Final Project', ''),
    ]
    for row_idx, (mod, topic, notes) in enumerate(modules, start=1):
        mod_table.cell(row_idx, 0).text = mod
        mod_table.cell(row_idx, 1).text = topic
        mod_table.cell(row_idx, 2).text = notes

    doc.add_paragraph()

    # Learning Goals
    p = doc.add_paragraph()
    p.add_run('My Learning Goals:').bold = True

    for i in range(1, 4):
        doc.add_paragraph(f'{i}. _________________________________')

    doc.add_paragraph()

    # ===== MODULE 1 =====
    add_page_break(doc)
    doc.add_heading('Module 1: Foundation - Google Workspace with AI', level=2)

    p = doc.add_paragraph()
    p.add_run('Session Objectives:').bold = True
    for obj in ['Navigate the Google Workspace ecosystem', 'Locate Gemini AI in core applications', 'Send first AI-assisted prompts']:
        doc.add_paragraph(f'• {obj}', style='List Bullet')

    # 5 Core Apps table
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('The 5 Core Apps:').bold = True

    core_table = doc.add_table(rows=6, cols=3)
    core_table.style = 'Table Grid'
    core_headers = ['App', 'My Purpose', 'Gemini Location']
    for i, h in enumerate(core_headers):
        core_table.cell(0, i).text = h
        core_table.cell(0, i).paragraphs[0].runs[0].bold = True

    apps = ['Gmail', 'Calendar', 'Drive', 'Docs', 'Sheets']
    for i, app in enumerate(apps, start=1):
        core_table.cell(i, 0).text = app

    doc.add_paragraph()

    # My AI Prompts
    p = doc.add_paragraph()
    p.add_run('My First AI Prompts:').bold = True

    doc.add_paragraph('Gmail Prompt: _________________________________')
    doc.add_paragraph('Docs Prompt: _________________________________')
    doc.add_paragraph('Sheets Prompt: _________________________________')

    doc.add_paragraph()
    doc.add_paragraph('Session Notes:')
    doc.add_paragraph('_________________________________')
    doc.add_paragraph('_________________________________')

    # Compression Artifact
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Module 1 Compression Artifact').bold = True
    p = doc.add_paragraph()
    p.add_run('Gemini Hotkeys Quick Reference').bold = True

    hotkey_table = doc.add_table(rows=7, cols=2)
    hotkey_table.style = 'Table Grid'
    hotkeys = [
        ('Gmail', '________________________________'),
        ('Docs', '________________________________'),
        ('Sheets', '________________________________'),
        ('Slides', '________________________________'),
        ('Drive', '________________________________'),
        ('Calendar', '________________________________'),
        ('Meet', '________________________________'),
    ]
    for i, (app, key) in enumerate(hotkeys):
        hotkey_table.cell(i, 0).text = app
        hotkey_table.cell(i, 1).text = key

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('My Commitment:').bold = True
    doc.add_paragraph('One thing I\'ll try before Module 2: _________________________________')

    doc.add_paragraph()
    doc.add_paragraph('─' * 60)
    doc.add_paragraph()

    # ===== MODULE 2 =====
    add_page_break(doc)
    doc.add_heading('Module 2: Collaboration with AI', level=2)

    p = doc.add_paragraph()
    p.add_run('Session Objectives:').bold = True
    for obj in ['Share documents for real-time collaboration', 'Use Suggesting mode for tracked changes', 'Generate and respond to comments with AI']:
        doc.add_paragraph(f'• {obj}', style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Collaboration Modes:').bold = True

    collab_table = doc.add_table(rows=3, cols=2)
    collab_table.style = 'Table Grid'
    collab_table.cell(0, 0).text = 'Mode'
    collab_table.cell(0, 1).text = 'When to Use'
    for cell in collab_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
    collab_table.cell(1, 0).text = 'Editing'
    collab_table.cell(2, 0).text = 'Suggesting'

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('My AI-Enhanced Comments:').bold = True
    doc.add_paragraph('Comment 1: _________________________________')
    doc.add_paragraph('AI Suggestion: _________________________________')
    doc.add_paragraph('My Response: _________________________________')

    doc.add_paragraph()
    doc.add_paragraph('Session Notes:')
    doc.add_paragraph('_________________________________')
    doc.add_paragraph('_________________________________')

    # Compression Artifact
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Module 2 Compression Artifact').bold = True
    p = doc.add_paragraph()
    p.add_run('AI Collaboration Checklist:').bold = True

    for item in ['Enabled Suggesting mode', 'Shared with "Anyone can edit"', 'Added AI comment', 'Replied using AI', 'Resolved thread']:
        doc.add_paragraph(f'☐ {item}')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('My Commitment:').bold = True
    doc.add_paragraph('One thing I\'ll try before Module 3: _________________________________')

    doc.add_paragraph()
    doc.add_paragraph('─' * 60)
    doc.add_paragraph()

    # ===== MODULE 3 =====
    add_page_break(doc)
    doc.add_heading('Module 3: Document Automation', level=2)

    p = doc.add_paragraph()
    p.add_run('Session Objectives:').bold = True
    for obj in ['Create reusable document templates', 'Fill templates using AI', 'Generate content variations']:
        doc.add_paragraph(f'• {obj}', style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('My Template Sections:').bold = True
    for i in range(1, 6):
        doc.add_paragraph(f'{i}. _________________________________')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('AI Variation Results:').bold = True
    doc.add_paragraph('Original: _________________________________')
    doc.add_paragraph('Variation 1 (Formal): _________________________________')
    doc.add_paragraph('Variation 2 (Casual): _________________________________')
    doc.add_paragraph('Variation 3 (Bullet): _________________________________')

    doc.add_paragraph()
    doc.add_paragraph('Session Notes:')
    doc.add_paragraph('_________________________________')
    doc.add_paragraph('_________________________________')

    # Template Starter Kit
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Module 3 Compression Artifact').bold = True
    p = doc.add_paragraph()
    p.add_run('My Template Starter Kit:').bold = True

    template_table = doc.add_table(rows=6, cols=2)
    template_table.style = 'Table Grid'
    template_table.cell(0, 0).text = 'Template'
    template_table.cell(0, 1).text = 'When I\'ll Use It'
    for cell in template_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
    for i in range(1, 6):
        template_table.cell(i, 0).text = f'{i}.'

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('My Commitment:').bold = True
    doc.add_paragraph('One thing I\'ll try before Module 4: _________________________________')

    doc.add_paragraph()
    doc.add_paragraph('─' * 60)
    doc.add_paragraph()

    # ===== MODULE 4 =====
    add_page_break(doc)
    doc.add_heading('Module 4: AI-Powered Research', level=2)

    p = doc.add_paragraph()
    p.add_run('Session Objectives:').bold = True
    for obj in ['Extract information from multiple sources', 'Synthesize findings using Gemini', 'Generate structured research outputs']:
        doc.add_paragraph(f'• {obj}', style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Research Synthesis:').bold = True
    doc.add_paragraph('Topic: _________________________________')
    doc.add_paragraph('Key Finding 1: _________________________________')
    doc.add_paragraph('Evidence: _________________________________')
    doc.add_paragraph('Key Finding 2: _________________________________')
    doc.add_paragraph('Evidence: _________________________________')
    doc.add_paragraph('Key Finding 3: _________________________________')
    doc.add_paragraph('Evidence: _________________________________')
    doc.add_paragraph('Source Conflicts Identified: _________________________________')

    doc.add_paragraph()
    doc.add_paragraph('Session Notes:')
    doc.add_paragraph('_________________________________')
    doc.add_paragraph('_________________________________')

    # Research Workflow
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Module 4 Compression Artifact').bold = True
    p = doc.add_paragraph()
    p.add_run('Research Workflow:').bold = True

    for item in ['Collect sources (3-5)', 'Extract key data', 'Synthesize with Gemini', 'Verify output', 'Add citations']:
        doc.add_paragraph(f'☐ {item}')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('My Commitment:').bold = True
    doc.add_paragraph('One thing I\'ll try before Module 5: _________________________________')

    doc.add_paragraph()
    doc.add_paragraph('─' * 60)
    doc.add_paragraph()

    # ===== MODULE 5 =====
    add_page_break(doc)
    doc.add_heading('Module 5: Meetings & Calendaring', level=2)

    p = doc.add_paragraph()
    p.add_run('Session Objectives:').bold = True
    for obj in ['Schedule meetings using AI assistance', 'Generate meeting agendas with Gemini', 'Draft meeting notes and follow-ups']:
        doc.add_paragraph(f'• {obj}', style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('My AI-Generated Agenda:').bold = True
    doc.add_paragraph('Meeting Type: _________________________________')
    doc.add_paragraph('Agenda Items:')
    for i in range(1, 5):
        doc.add_paragraph(f'{i}. _________________________________')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Post-Meeting AI Extraction:').bold = True
    doc.add_paragraph('Decisions Made:')
    doc.add_paragraph('1. _________________________________')
    doc.add_paragraph('2. _________________________________')
    doc.add_paragraph('Action Items:')

    action_table = doc.add_table(rows=3, cols=3)
    action_table.style = 'Table Grid'
    action_table.cell(0, 0).text = 'Action'
    action_table.cell(0, 1).text = 'Owner'
    action_table.cell(0, 2).text = 'Due Date'
    for cell in action_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph('Session Notes:')
    doc.add_paragraph('_________________________________')
    doc.add_paragraph('_________________________________')

    # Meeting Automation
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Module 5 Compression Artifact').bold = True
    p = doc.add_paragraph()
    p.add_run('Meeting Automation Script:').bold = True

    for item in ['Schedule using AI', 'Generate agenda', 'Prepare materials', 'During: AI notes', 'Post-meeting: Extract & summarize', 'Create tasks', 'Archive notes']:
        doc.add_paragraph(f'☐ {item}')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('My Commitment:').bold = True
    doc.add_paragraph('One thing I\'ll try before Module 6: _________________________________')

    doc.add_paragraph()
    doc.add_paragraph('─' * 60)
    doc.add_paragraph()

    # ===== MODULE 6 =====
    add_page_break(doc)
    doc.add_heading('Module 6: Drive & File Management', level=2)

    p = doc.add_paragraph()
    p.add_run('Session Objectives:').bold = True
    for obj in ['Use natural language search to find files', 'Create systematic folder structures', 'Set appropriate sharing permissions']:
        doc.add_paragraph(f'• {obj}', style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('My AI Searches route:').bold = True

    search_table = doc.add_table(rows=4, cols=3)
    search_table.style = 'Table Grid'
    search_table.cell(0, 0).text = 'Search route'
    search_table.cell(0, 1).text = 'Results'
    search_table.cell(0, 2).text = 'Accuracy'
    for cell in search_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('My Folder Structure:').bold = True

    # Folder structure display
    folder_para = doc.add_paragraph()
    folder_para.add_run('[Project Name]/').bold = True
    for i in range(1, 6):
        doc.add_paragraph(f'├── ________________________________')

    doc.add_paragraph()
    doc.add_paragraph('Session Notes:')
    doc.add_paragraph('_________________________________')
    doc.add_paragraph('_________________________________')

    # Drive Organization
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Module 6 Compression Artifact').bold = True
    p = doc.add_paragraph()
    p.add_run('Drive Organization System:').bold = True

    doc.add_paragraph('My Current Organization Problem: _________________________________')
    doc.add_paragraph('My New System:')
    for item in ['Number folders', 'Date format: YYYY-MM-DD', 'Star priority files', 'Monthly review scheduled']:
        doc.add_paragraph(f'☐ {item}')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('My Commitment:').bold = True
    doc.add_paragraph('One thing I\'ll try before Module 7: _________________________________')

    doc.add_paragraph()
    doc.add_paragraph('─' * 60)
    doc.add_paragraph()

    # ===== MODULE 7 =====
    add_page_break(doc)
    doc.add_heading('Module 7: Automation with Gemini', level=2)

    p = doc.add_paragraph()
    p.add_run('Session Objectives:').bold = True
    for obj in ['Identify automation opportunities', 'Build prompt chains', 'Create personal automation playbook']:
        doc.add_paragraph(f'• {obj}', style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('My Repetitive Tasks:').bold = True

    task_table = doc.add_table(rows=4, cols=4)
    task_table.style = 'Table Grid'
    task_table.cell(0, 0).text = 'Task'
    task_table.cell(0, 1).text = 'Frequency'
    task_table.cell(0, 2).text = 'Time/Instance'
    task_table.cell(0, 3).text = 'Total Time'
    for cell in task_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('My Prompt Chain:').bold = True
    doc.add_paragraph('Automation Name: _________________________________')

    for step in range(1, 4):
        doc.add_paragraph(f'Step {step}:')
        doc.add_paragraph('Prompt: _________________________________')
        doc.add_paragraph('Output: _________________________________')

    doc.add_paragraph()
    doc.add_paragraph('Session Notes:')
    doc.add_paragraph('_________________________________')
    doc.add_paragraph('_________________________________')

    # Top 3 Patterns
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Module 7 Compression Artifact').bold = True
    p = doc.add_paragraph()
    p.add_run('My Top 3 Automation Patterns:').bold = True

    for i in range(1, 4):
        doc.add_paragraph(f'Pattern {i}: _________________________________')
        doc.add_paragraph(f'Used for: _________________________________')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('My Commitment:').bold = True
    doc.add_paragraph('One automation I\'ll build this week: _________________________________')

    doc.add_paragraph()
    doc.add_paragraph('─' * 60)
    doc.add_paragraph()

    # ===== MODULE 8 =====
    add_page_break(doc)
    doc.add_heading('Module 8: AI Integration Final Project', level=2)

    p = doc.add_paragraph()
    p.add_run('My AI Workspace Playbook').bold = True
    doc.add_paragraph('Document Title: [Your Name] AI Workspace Playbook')

    # 5 Workflows
    for wf_num in range(1, 6):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(f'Workflow {wf_num}: [Name]').bold = True
        doc.add_paragraph('Trigger: _________________________________')
        doc.add_paragraph('Apps Used: _________________________________')
        doc.add_paragraph('Steps:')
        for i in range(1, 6):
            doc.add_paragraph(f'{i}. _________________________________')
        doc.add_paragraph('Time Saved: _________________________________')

        p = doc.add_paragraph()
        p.add_run('Gemini Prompts Used:').bold = True
        doc.add_paragraph('1. _________________________________')
        doc.add_paragraph('2. _________________________________')
        doc.add_paragraph('3. _________________________________')

        if wf_num < 5:
            doc.add_paragraph()
            doc.add_paragraph('─' * 40)

    doc.add_paragraph()
    doc.add_paragraph('─' * 60)
    doc.add_paragraph()

    # ===== COURSE REFLECTION =====
    add_page_break(doc)
    doc.add_heading('Course Reflection', level=2)

    p = doc.add_paragraph()
    p.add_run('Biggest Breakthrough:').bold = True
    doc.add_paragraph('_________________________________')
    doc.add_paragraph('_________________________________')

    p = doc.add_paragraph()
    p.add_run('Most Challenging Module:').bold = True
    doc.add_paragraph('Module: _________________________________')
    doc.add_paragraph('Why: _________________________________')
    doc.add_paragraph('_________________________________')

    p = doc.add_paragraph()
    p.add_run('How I\'ll Continue Learning:').bold = True
    doc.add_paragraph('_________________________________')
    doc.add_paragraph('_________________________________')

    p = doc.add_paragraph()
    p.add_run('My Commitment to Continued Practice:').bold = True
    doc.add_paragraph('This week: I will try _________________________________')
    doc.add_paragraph('This month: I will _________________________________')
    doc.add_paragraph('This quarter: I will _________________________________')

    doc.add_paragraph()
    doc.add_paragraph('─' * 60)
    doc.add_paragraph()

    # ===== NOTES SECTION =====
    add_page_break(doc)
    doc.add_heading('Notes Section', level=2)

    for i in range(20):
        doc.add_paragraph('_________________________________')

    doc.add_paragraph()
    doc.add_paragraph('─' * 60)
    doc.add_paragraph()

    # ===== FINAL MESSAGE =====
    p = doc.add_paragraph()
    p.add_run('Thank you for completing the Google Workspace with AI Microlearning Course!').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    doc.save(output_path)
    print(f'Workbook saved to: {output_path}')

if __name__ == '__main__':
    output = '/home/tp/Documents/app-docs/micro-google-workplace/draft-1/participant-workbook/participant-workbook.docx'
    create_workbook_docx(output)